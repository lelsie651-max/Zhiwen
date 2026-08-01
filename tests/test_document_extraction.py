import asyncio
import io
import inspect
import tempfile
from pathlib import Path

import pytest
from docx import Document as DocxDocumentFactory
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from app.schemas.document_extraction import ExtractedBlockType, ExtractionOutcome
from app.schemas.file_ingestion import DetectedFileFormat
from app.services import document_extraction
from app.services.document_extraction import extract_document
from app.storage import LocalFileStorage
from app.utils import build_document_block_anchor_hash


def run_async(awaitable):
    return asyncio.run(awaitable)


def make_storage(tmp_path: Path) -> LocalFileStorage:
    return LocalFileStorage(
        temp_root=tmp_path / "tmp",
        storage_root=tmp_path / "files",
    )


async def store_bytes(storage: LocalFileStorage, data: bytes, *, suffix: str) -> str:
    temp_file = await storage.create_temp_file(suffix=suffix)
    await storage.write_temp_chunk(temp_file, data)
    return await storage.promote_temp_file(temp_file, suffix=suffix)


def build_pdf_with_text(pages: list[str | None]) -> bytes:
    writer = PdfWriter()
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)

    for page_text in pages:
        page = writer.add_blank_page(width=300, height=300)
        page[NameObject("/Resources")] = DictionaryObject(
            {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
        )
        if page_text:
            lines = page_text.split("\n")
            commands = ["BT", "/F1 12 Tf", "20 260 Td"]
            for index, line in enumerate(lines):
                escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
                if index > 0:
                    commands.append("0 -16 Td")
                commands.append(f"({escaped}) Tj")
            commands.append("ET")
            stream = DecodedStreamObject()
            stream.set_data("\n".join(commands).encode("utf-8"))
            page[NameObject("/Contents")] = writer._add_object(stream)

    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def build_docx_bytes() -> bytes:
    document = DocxDocumentFactory()
    document.add_heading("总标题", level=1)
    document.add_paragraph("第一段正文")
    list_paragraph = document.add_paragraph("列表项一")
    list_paragraph.style = "List Bullet"
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "列1"
    table.rows[0].cells[1].text = "列2"
    table.rows[1].cells[0].text = "值A"
    table.rows[1].cells[1].text = "值B"
    document.add_heading("子标题", level=2)
    document.add_paragraph("最后一段")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_pdf_text_is_extracted_by_page_with_real_page_numbers(tmp_path: Path) -> None:
    storage = make_storage(tmp_path)
    storage_key = run_async(
        store_bytes(
            storage,
            build_pdf_with_text(["第一页第一段\n\n第一页第二段", "第二页正文"]),
            suffix=".pdf",
        )
    )

    result = run_async(extract_document(storage, storage_key, DetectedFileFormat.PDF))

    assert result.outcome == ExtractionOutcome.SUCCESS
    assert result.page_count == 2
    assert [block.page_no for block in result.blocks] == [1, 2]
    assert all(block.block_type == ExtractedBlockType.PAGE_TEXT for block in result.blocks)


def test_pdf_without_text_returns_needs_ocr(tmp_path: Path) -> None:
    storage = make_storage(tmp_path)
    storage_key = run_async(
        store_bytes(
            storage,
            build_pdf_with_text([None]),
            suffix=".pdf",
        )
    )

    result = run_async(extract_document(storage, storage_key, DetectedFileFormat.PDF))

    assert result.outcome == ExtractionOutcome.NEEDS_OCR
    assert result.block_count == 0


def test_docx_keeps_paragraph_heading_list_and_table_order(tmp_path: Path) -> None:
    storage = make_storage(tmp_path)
    storage_key = run_async(store_bytes(storage, build_docx_bytes(), suffix=".docx"))

    result = run_async(extract_document(storage, storage_key, DetectedFileFormat.DOCX))

    assert result.outcome == ExtractionOutcome.SUCCESS
    assert [block.block_type for block in result.blocks] == [
        ExtractedBlockType.HEADING,
        ExtractedBlockType.PARAGRAPH,
        ExtractedBlockType.LIST_ITEM,
        ExtractedBlockType.TABLE_ROW,
        ExtractedBlockType.TABLE_ROW,
        ExtractedBlockType.HEADING,
        ExtractedBlockType.PARAGRAPH,
    ]
    assert result.blocks[0].heading_level == 1
    assert result.blocks[0].heading_path == ["总标题"]
    assert result.blocks[3].metadata["cells"] == ["列1", "列2"]
    assert result.blocks[5].heading_path == ["总标题", "子标题"]


def test_docx_does_not_invent_page_numbers(tmp_path: Path) -> None:
    storage = make_storage(tmp_path)
    storage_key = run_async(store_bytes(storage, build_docx_bytes(), suffix=".docx"))

    result = run_async(extract_document(storage, storage_key, DetectedFileFormat.DOCX))

    assert all(block.page_no is None for block in result.blocks)


def test_docx_heading_uses_style_id_even_when_display_name_is_localized(
    tmp_path: Path,
    monkeypatch,
) -> None:
    storage = make_storage(tmp_path)
    storage_key = run_async(store_bytes(storage, build_docx_bytes(), suffix=".docx"))
    original_get_style_name = document_extraction._get_docx_style_name

    def fake_get_style_name(style) -> str:
        style_id = document_extraction._get_docx_style_id(style)
        if style_id.lower() == "heading1":
            return "标题一"
        return original_get_style_name(style)

    monkeypatch.setattr(document_extraction, "_get_docx_style_name", fake_get_style_name)

    result = run_async(extract_document(storage, storage_key, DetectedFileFormat.DOCX))

    assert result.blocks[0].block_type == ExtractedBlockType.HEADING
    assert result.blocks[0].heading_level == 1
    assert result.blocks[0].heading_path == ["总标题"]
    assert result.blocks[5].heading_path == ["总标题", "子标题"]


def test_markdown_heading_paths_code_table_and_line_numbers(tmp_path: Path) -> None:
    storage = make_storage(tmp_path)
    markdown_text = "\n".join(
        [
            "# Title",
            "Intro line 1",
            "Intro line 2",
            "",
            "- item one",
            "```python",
            "print('hi')",
            "```",
            "| Name | Value |",
            "| --- | --- |",
            "| A | 1 |",
            "## Child",
            "Tail paragraph",
        ]
    )
    storage_key = run_async(store_bytes(storage, markdown_text.encode("utf-8"), suffix=".md"))

    result = run_async(extract_document(storage, storage_key, DetectedFileFormat.MD))

    assert result.outcome == ExtractionOutcome.SUCCESS
    assert result.blocks[0].heading_path == ["Title"]
    assert result.blocks[1].start_line == 2 and result.blocks[1].end_line == 3
    assert result.blocks[2].block_type == ExtractedBlockType.LIST_ITEM
    assert result.blocks[3].block_type == ExtractedBlockType.CODE
    assert result.blocks[3].start_line == 6 and result.blocks[3].end_line == 8
    table_rows = [block for block in result.blocks if block.block_type == ExtractedBlockType.TABLE_ROW]
    assert len(table_rows) == 2
    assert table_rows[0].table_index == 1
    assert table_rows[1].table_index == 1
    assert table_rows[0].row_index == 1
    assert table_rows[1].row_index == 2
    assert table_rows[0].location_key == "md:t1:r1:l9"
    assert table_rows[1].location_key == "md:t1:r2:l11"
    assert table_rows[0].start_line == 9
    assert table_rows[1].start_line == 11
    assert result.blocks[-1].heading_path == ["Title", "Child"]


def test_markdown_second_table_gets_incremented_table_index(tmp_path: Path) -> None:
    storage = make_storage(tmp_path)
    markdown_text = "\n".join(
        [
            "| A | B |",
            "| --- | --- |",
            "| 1 | 2 |",
            "",
            "Paragraph break",
            "",
            "| X | Y |",
            "| --- | --- |",
            "| 3 | 4 |",
        ]
    )
    storage_key = run_async(store_bytes(storage, markdown_text.encode("utf-8"), suffix=".md"))

    result = run_async(extract_document(storage, storage_key, DetectedFileFormat.MD))

    table_rows = [block for block in result.blocks if block.block_type == ExtractedBlockType.TABLE_ROW]
    assert [block.table_index for block in table_rows] == [1, 1, 2, 2]
    assert [block.row_index for block in table_rows] == [1, 2, 1, 2]


def test_markdown_four_backticks_are_not_closed_by_three(tmp_path: Path) -> None:
    storage = make_storage(tmp_path)
    markdown_text = "\n".join(
        [
            "````python",
            "print('hi')",
            "```",
            "still code",
            "````",
        ]
    )
    storage_key = run_async(store_bytes(storage, markdown_text.encode("utf-8"), suffix=".md"))

    result = run_async(extract_document(storage, storage_key, DetectedFileFormat.MD))

    assert len(result.blocks) == 1
    assert result.blocks[0].block_type == ExtractedBlockType.CODE
    assert "still code" in result.blocks[0].raw_text
    assert result.blocks[0].end_line == 5


def test_markdown_tilde_fence_is_recognized(tmp_path: Path) -> None:
    storage = make_storage(tmp_path)
    markdown_text = "\n".join(
        [
            "~~~text",
            "hello",
            "~~~",
        ]
    )
    storage_key = run_async(store_bytes(storage, markdown_text.encode("utf-8"), suffix=".md"))

    result = run_async(extract_document(storage, storage_key, DetectedFileFormat.MD))

    assert len(result.blocks) == 1
    assert result.blocks[0].block_type == ExtractedBlockType.CODE
    assert result.blocks[0].raw_text.startswith("~~~text")


def test_markdown_unclosed_fence_emits_warning_and_keeps_code_block(tmp_path: Path) -> None:
    storage = make_storage(tmp_path)
    markdown_text = "\n".join(
        [
            "```python",
            "print('hi')",
            "| not | table |",
            "# not heading",
        ]
    )
    storage_key = run_async(store_bytes(storage, markdown_text.encode("utf-8"), suffix=".md"))

    result = run_async(extract_document(storage, storage_key, DetectedFileFormat.MD))

    assert len(result.blocks) == 1
    assert result.blocks[0].block_type == ExtractedBlockType.CODE
    assert "| not | table |" in result.blocks[0].raw_text
    assert "# not heading" in result.blocks[0].raw_text
    assert result.warnings == ["Markdown 存在未闭合代码围栏。"]


def test_txt_uses_encoding_hint_and_preserves_line_numbers(tmp_path: Path) -> None:
    storage = make_storage(tmp_path)
    txt_text = "第一段\n第二行\n\n第二段"
    storage_key = run_async(store_bytes(storage, txt_text.encode("utf-16"), suffix=".txt"))

    result = run_async(
        extract_document(
            storage,
            storage_key,
            DetectedFileFormat.TXT,
            encoding_hint="utf-16",
        )
    )

    assert result.outcome == ExtractionOutcome.SUCCESS
    assert result.detected_encoding == "utf-16"
    assert result.blocks[0].start_line == 1 and result.blocks[0].end_line == 2
    assert result.blocks[1].start_line == 4 and result.blocks[1].end_line == 4


def test_source_order_is_continuous_and_location_keys_are_unique(tmp_path: Path) -> None:
    storage = make_storage(tmp_path)
    markdown_text = "# T\n\nPara one\n\nPara two"
    storage_key = run_async(store_bytes(storage, markdown_text.encode("utf-8"), suffix=".md"))

    result = run_async(extract_document(storage, storage_key, DetectedFileFormat.MD))

    assert [block.source_order for block in result.blocks] == list(range(len(result.blocks)))
    assert len({block.location_key for block in result.blocks}) == len(result.blocks)


def test_same_input_produces_same_anchor_hashes(tmp_path: Path) -> None:
    storage = make_storage(tmp_path)
    content = "# 标题\n\n内容"
    storage_key = run_async(store_bytes(storage, content.encode("utf-8"), suffix=".md"))

    first = run_async(extract_document(storage, storage_key, DetectedFileFormat.MD))
    second = run_async(extract_document(storage, storage_key, DetectedFileFormat.MD))

    assert [block.anchor_hash for block in first.blocks] == [block.anchor_hash for block in second.blocks]


def test_document_extraction_uses_public_anchor_hash_function() -> None:
    source = inspect.getsource(document_extraction)
    assert document_extraction.build_document_block_anchor_hash is build_document_block_anchor_hash
    assert "def _build_anchor_hash" not in source


def test_long_block_is_safely_split(tmp_path: Path) -> None:
    storage = make_storage(tmp_path)
    long_text = "A" * 21000 + " " + "B" * 21000
    storage_key = run_async(store_bytes(storage, long_text.encode("utf-8"), suffix=".txt"))

    result = run_async(extract_document(storage, storage_key, DetectedFileFormat.TXT))

    assert result.block_count >= 2
    assert all(len(block.raw_text) <= 20000 for block in result.blocks)
    assert all("segment_count" in block.metadata for block in result.blocks)


@pytest.mark.parametrize("payload", [b"plain text", b"%PDF-bad"])
def test_spooled_temporary_file_is_always_closed(tmp_path: Path, monkeypatch, payload: bytes) -> None:
    storage = make_storage(tmp_path)
    suffix = ".txt" if payload == b"plain text" else ".pdf"
    detected_format = DetectedFileFormat.TXT if suffix == ".txt" else DetectedFileFormat.PDF
    storage_key = run_async(store_bytes(storage, payload, suffix=suffix))
    closed_states: list[bool] = []
    real_factory = tempfile.SpooledTemporaryFile

    class TrackingSpooledFile:
        def __init__(self, *args, **kwargs) -> None:
            self._inner = real_factory(*args, **kwargs)
            self.closed = False

        def write(self, data: bytes) -> int:
            return self._inner.write(data)

        def seek(self, offset: int, whence: int = 0) -> int:
            return self._inner.seek(offset, whence)

        def read(self, size: int = -1) -> bytes:
            return self._inner.read(size)

        def close(self) -> None:
            self._inner.close()
            self.closed = True
            closed_states.append(self.closed)

    monkeypatch.setattr(document_extraction.tempfile, "SpooledTemporaryFile", TrackingSpooledFile)

    run_async(extract_document(storage, storage_key, detected_format))

    assert closed_states == [True]


def test_failure_log_does_not_leak_storage_key(tmp_path: Path, monkeypatch, caplog) -> None:
    storage = make_storage(tmp_path)
    storage_key = run_async(store_bytes(storage, b"hello", suffix=".txt"))

    async def broken_reader(_storage_key: str, *, chunk_size: int):
        raise RuntimeError(f"boom {storage_key}")
        yield b""

    monkeypatch.setattr(storage, "read_file_chunks", broken_reader)

    with caplog.at_level("WARNING"):
        result = run_async(extract_document(storage, storage_key, DetectedFileFormat.TXT))

    assert result.outcome == ExtractionOutcome.FAILED
    assert storage_key not in caplog.text
